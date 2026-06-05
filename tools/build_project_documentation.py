from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "Potato_Disease_Classification_Project_Documentation.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B7C9DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value)
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_steps(document, steps):
    for step in steps:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(step)


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="D8E2ED")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title_run.font.size = Pt(10.5)
    p.add_run("\n")
    body_run = p.add_run(body)
    body_run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()


def set_document_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_document():
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Potato Disease Classification")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub = subtitle.add_run(
        "End-to-end deep learning project: model training, FastAPI backend, React frontend, and free deployment."
    )
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(80, 80, 80)

    add_callout(
        doc,
        "Live project summary",
        "This app lets a user upload a potato leaf image and returns a disease prediction with confidence. "
        "The model predicts Early Blight, Late Blight, or Healthy.",
    )

    add_heading(doc, "1. Problem Statement", 1)
    doc.add_paragraph(
        "The goal of this project is to identify potato leaf health from an image. "
        "A user uploads a leaf photo, and the application classifies it into one of three categories."
    )
    add_bullets(doc, ["Early Blight", "Late Blight", "Healthy"])

    add_heading(doc, "2. Data Collection and Preprocessing", 1)
    doc.add_paragraph(
        "The model was trained using potato leaf images organized by disease category. "
        "Images were prepared for model input by resizing them to a consistent shape and normalizing pixel values."
    )
    add_table(
        doc,
        ["Step", "What happened"],
        [
            ["Collect images", "Used categorized potato leaf images for Early Blight, Late Blight, and Healthy leaves."],
            ["Preprocess", "Converted images into arrays and resized them to the input size expected by the model."],
            ["Prepare labels", "Mapped each image to the correct class name for supervised learning."],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "3. Model Building", 1)
    doc.add_paragraph(
        "A Keras deep learning model was trained to learn visual patterns from potato leaf images. "
        "The final trained model was saved as a .keras file and used by the backend for live prediction."
    )
    add_bullets(
        doc,
        [
            "Input: potato leaf image",
            "Model file: saved_models/1/1.keras",
            "Output: predicted class and confidence score",
        ],
    )

    add_heading(doc, "4. FastAPI Backend", 1)
    doc.add_paragraph(
        "The backend was built with FastAPI. It exposes a REST endpoint that accepts an uploaded image, "
        "loads the Keras model, runs prediction, and returns a JSON response."
    )
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["/", "Health check endpoint. Returns a message confirming the API is running."],
            ["/prediction", "POST endpoint used by the React frontend to upload a leaf image and receive prediction results."],
        ],
        [1.8, 4.7],
    )
    add_callout(
        doc,
        "API URL",
        "https://vaibhavnayak-potato-disease-api.hf.space/prediction",
    )

    add_heading(doc, "5. React Website", 1)
    doc.add_paragraph(
        "The frontend was built in React. It provides a simple image upload interface, sends the selected image "
        "to the FastAPI endpoint, and displays the predicted label and confidence percentage."
    )
    add_bullets(
        doc,
        [
            "User uploads or drags a potato leaf image.",
            "React sends the image to the backend using REACT_APP_API_URL.",
            "The result is shown in a small table with label and confidence.",
        ],
    )

    add_heading(doc, "6. Free Deployment", 1)
    doc.add_paragraph(
        "The application was deployed using a free, beginner-friendly setup: Hugging Face Spaces for the ML API "
        "and Vercel for the React frontend."
    )
    add_table(
        doc,
        ["Part", "Platform", "Why it was used"],
        [
            ["Backend API", "Hugging Face Spaces", "Good free option for ML demos and Docker-based FastAPI apps."],
            ["Frontend", "Vercel", "Simple free hosting for React applications with GitHub-based deployment."],
            ["Model serving", "Docker", "Made the backend reproducible by defining Python, dependencies, and startup command."],
        ],
        [1.45, 1.75, 3.3],
    )

    add_heading(doc, "7. End-to-End Flow", 1)
    add_steps(
        doc,
        [
            "Train the potato leaf disease classification model in Keras.",
            "Save the trained model as a .keras file.",
            "Build a FastAPI backend and load the saved Keras model.",
            "Create a /prediction endpoint that accepts uploaded leaf images.",
            "Build a React frontend for image upload and result display.",
            "Deploy the backend to Hugging Face Spaces using Docker.",
            "Deploy the frontend to Vercel and connect it to the backend API URL.",
            "Share the Vercel app link so others can test it from the browser.",
        ],
    )

    add_heading(doc, "8. How To Test", 1)
    doc.add_paragraph(
        "I will try attaching a few leaf images along with the LinkedIn post so you can test the app quickly. "
        "If you do not find sample images attached, please navigate to the GitHub repository and feel free to use the sample leaf images from there."
    )
    add_callout(
        doc,
        "Reader note",
        "This project is for learning and demonstration. Real agricultural decisions should use expert validation and field-level diagnosis.",
    )

    add_heading(doc, "9. Credits", 1)
    doc.add_paragraph(
        "Thanks to Codebasics and Dhaval Patel for the learning resources and project inspiration that helped shape this end-to-end machine learning application."
    )

    add_heading(doc, "10. Tech Stack", 1)
    add_table(
        doc,
        ["Layer", "Tools"],
        [
            ["Model", "Python, TensorFlow, Keras"],
            ["Backend", "FastAPI, Uvicorn, Pillow, NumPy"],
            ["Frontend", "React, Axios, Material UI"],
            ["Deployment", "Hugging Face Spaces, Docker, Vercel"],
        ],
        [1.8, 4.7],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Potato Disease Classification | Built by Vaibhav Nayak")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
